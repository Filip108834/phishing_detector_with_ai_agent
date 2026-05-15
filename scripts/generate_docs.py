"""Generuje plik Word z historią projektu."""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(11)


def h1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

def h2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

def h3(text):
    doc.add_heading(text, level=3)

def para(text):
    doc.add_paragraph(text)

def bullet(text):
    doc.add_paragraph(text, style='List Bullet')

def code(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), 'F3F4F6')
    shd.set(qn('w:val'), 'clear')
    p._p.get_or_add_pPr().append(shd)

def table_with_header(headers, rows_data):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
    for row_data in rows_data:
        row = t.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = val
    doc.add_paragraph()


# ---- TYTUŁ ----
doc.add_paragraph()
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('Agent AI do detekcji phishingu')
r.bold = True
r.font.size = Pt(22)
r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

t2 = doc.add_paragraph()
t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
t2.add_run('Historia i dokumentacja techniczna projektu\n'
           'Autor: Filip Pich | Praca inzynierska 2025\n'
           'Python 3.11 | Docker | FastAPI | scikit-learn | spaCy | Gophish')
doc.add_page_break()

# ---- 1. OPIS PROJEKTU ----
h1('1. Opis projektu')
para('Celem pracy inzynierskiej jest opracowanie i implementacja systemu umozliwiajacego '
     'symulacje atakow phishingowych oraz ocene skutecznosci mechanizmow obronnych z '
     'wykorzystaniem agenta sztucznej inteligencji.')
para('Zakres projektu:')
bullet('Srodowisko testowe do kampanii phishingowych (Gophish + MailHog)')
bullet('Agent AI analizujacy tresc wiadomosci e-mail, linki i metadane')
bullet('Klasyfikacja metodami ML: TF-IDF, modele klasyfikacyjne, NLP (spaCy, NLTK)')
bullet('Badanie skutecznosci i formulowanie wnioskow')
doc.add_page_break()

# ---- 2. PLAN TECHNICZNY ----
h1('2. Plan techniczny i stack technologiczny')

h2('2.1 Architektura systemu (4 warstwy)')
bullet('Warstwa symulacji: Gophish -> MailHog')
bullet('Warstwa ingestion: Parser e-maili (naglowki, body, URL-e)')
bullet('Agent AI: Feature Extraction -> ML Pipeline -> Klasyfikacja')
bullet('Warstwa storage + API: PostgreSQL <-> FastAPI <-> Modul metryk')

h2('2.2 Stack technologiczny')
h3('Jezyk i runtime')
bullet('Python 3.11, Docker + Docker Compose, Visual Studio Code')
h3('API agenta')
bullet('FastAPI + uvicorn, pydantic v2')
h3('Machine Learning i NLP')
bullet('scikit-learn: Pipeline, ColumnTransformer, VotingClassifier, GridSearchCV')
bullet('pandas, numpy, scipy')
bullet('spaCy (en_core_web_sm) - NER, POS tagging, segmentacja zdan')
bullet('NLTK - stemming (SnowballStemmer), stopwords, tokenizacja')
h3('Parsowanie e-maili i analiza URL')
bullet('email (stdlib), beautifulsoup4, lxml, tldextract, requests')
h3('Symulacja phishingu')
bullet('Gophish - silnik kampanii phishingowych')
bullet('MailHog - testowy serwer SMTP/inbox')
h3('Baza danych')
bullet('PostgreSQL 16 (glowna), SQLAlchemy ORM')

h2('2.3 Fazy implementacji')
bullet('Faza 1 - Srodowisko i dane: Docker, Gophish, MailHog, datasety publiczne')
bullet('Faza 2 - Pipeline ML: parser e-maili, feature engineering, Logistic Regression + Random Forest')
bullet('Faza 3 - Agent AI rozszerzony: integracja NLP (spaCy), NLTK preprocessor, EDA notebook')
bullet('Faza 4 - API i integracja: refaktoryzacja routerow, historia predykcji, MailHog poller')
bullet('Faza 5 - Ewaluacja: metryki, eksperymenty, raporty')

h2('2.4 Dane treningowe')
table_with_header(
    ['Dataset', 'Opis', 'Zrodlo'],
    [
        ('Enron Email Dataset', '~517k e-maili korporacyjnych (legit)', 'CMU / Kaggle'),
        ('SpamAssassin Public Corpus', 'spam + ham, ~3700 wiadomosci', 'Apache'),
        ('Dane wlasne (Gophish)', 'E-maile z kampanii testowych (labeled)', 'Lab'),
    ]
)
doc.add_page_break()

# ---- 3. FAZA 1 ----
h1('3. Faza 1 - Srodowisko laboratoryjne i dane')

h2('3.1 Stworzone pliki')
bullet('.env.example - szablon zmiennych srodowiskowych')
bullet('docker/gophish/config.json - konfiguracja Gophish (TLS wylaczone)')
bullet('docker-compose.yml - zaktualizowany: volume gophish_data, env_file')
bullet('simulation/gophish_client.py - wrapper Gophish REST API v1')
bullet('simulation/mailhog_client.py - klient MailHog API (fetch, eksport .eml)')
bullet('simulation/campaign_generator.py - 3 szablony phishing + 2 legit -> data/raw/')
bullet('scripts/fetch_datasets.py - pobieranie SpamAssassin + instrukcja Enron')

h2('3.2 Szablony e-mail')
bullet('lab_phish_bank_alert - alert o zablokowaniu konta bankowego')
bullet('lab_phish_password_reset - wygasniecie hasla Microsoft')
bullet('lab_phish_package_delivery - doplata za przesylke kurierska')
bullet('lab_legit_newsletter - miesieczny newsletter dzialu IT')
bullet('lab_legit_meeting_invite - zaproszenie na spotkanie')

h2('3.3 Siec Docker')
bullet('lab_internal (bridge, internal=true): Gophish <-> MailHog <-> AI Agent')
bullet('db_internal (bridge, internal=true): AI Agent <-> PostgreSQL')
bullet('Wszystkie porty bindowane na 127.0.0.1')

h2('3.4 Workflow')
code('cp .env.example .env\ndocker compose up -d\ndocker compose logs gophish | grep "Please login"\npython scripts/fetch_datasets.py\npython -m simulation.campaign_generator --mode all')
doc.add_page_break()

# ---- 4. REFACTORING ----
h1('4. Refactoring strukturalny repozytorium')

h2('4.1 Usuniete')
bullet('proejct-root/ - katalog z literowka w nazwie (duplikat agent/main.py), usuniety z git')

h2('4.2 Naprawione bledy .gitignore')
table_with_header(
    ['Problem', 'Bylo', 'Jest'],
    [
        ('.env.example ignorowany', 'brak wyjatku', '!.env.example'),
        ('.vscode/settings.json blokowany', '.vscode/ (katalog)', '.vscode/* + !.vscode/settings.json'),
        ('data/raw/ ignorowalo .gitkeep', 'data/raw/', 'data/raw/** + !data/raw/.gitkeep'),
        ('models/ ignorowalo .gitkeep', 'models/', 'models/** + !models/.gitkeep'),
    ]
)

h2('4.3 Poprawiony .dockerignore')
bullet('Dodano: data/datasets/, notebooks/, experiments/, docs/, .git/, *.md, LICENSE')
bullet('Obraz produkcyjny zawiera tylko runtime')
doc.add_page_break()

# ---- 5. FAZA 2 ----
h1('5. Faza 2 - Pipeline ML (podstawowy)')

h2('5.1 agent/ingestion/')
bullet('email_parser.py - ParsedEmail z MIME: naglowki, plain text + HTML (BeautifulSoup), '
       'URL-e, wynik SPF, liczba przeskokow Received. Fabryki: from_string(), from_file(), from_request()')
bullet('dataset_loader.py - laduje SpamAssassin, kampanie Gophish i Enron CSV '
       'do DataFrame {raw_eml, label, source}')

h2('5.2 agent/features/')
bullet('url_analyzer.py - 11 cech URL: IP, podejrzane TLD, slowa kluczowe, @ w URL, '
       'redirect param, dlugosci, glebokosc subdomen, entropia Shannona')
bullet('extractor.py - 23 cechy numeryczne + pole text dla TF-IDF')

h2('5.3 agent/ml/pipeline.py')
bullet('ColumnTransformer: TF-IDF(text, ngram(1,2), 8000) + StandardScaler(feat_*)')
bullet('VotingClassifier(LR + RF + GB, soft voting)')
bullet('Funkcje: build_pipeline(), train(), cross_validate_pipeline(), evaluate(), save()/load()')

h2('5.4 experiments/train_model.py')
para('CLI: ladowanie danych -> ekstrakcja cech -> train/test split (80/20, stratified) '
     '-> trening -> CV (5-fold) -> metryki -> zapis modelu i wynikow JSON.')
code('python -m experiments.train_model --compare --max-per-class 1000\npython -m experiments.train_model --model ensemble')
doc.add_page_break()

# ---- 6. FAZA 3 ----
h1('6. Faza 3 - Agent AI (rozszerzony NLP)')

h2('6.1 agent/features/nlp_analyzer.py (spaCy)')
para('Model en_core_web_sm ladowany leniwie. Fallback do 0.0 gdy model niedostepny.')
para('12 nowych cech (prefix feat_nlp_):')
bullet('NER: ner_org_count, ner_money_count, ner_date_count, ner_cardinal_count')
bullet('POS: verb_density, noun_density, function_word_ratio, imperative_count')
bullet('Zdania: sentence_count, avg_sentence_len, exclaim_sentence_ratio')
bullet('Leksykalne: unique_token_ratio (TTR)')

h2('6.2 agent/features/text_preprocessor.py (NLTK)')
para('sklearn transformer: lowercase -> normalizacja URL/email/liczb -> tokenizacja '
     '-> usuwanie EN+PL stopwords -> stemming (SnowballStemmer). '
     'Zasoby NLTK pobierane automatycznie.')

h2('6.3 Zmiany')
bullet('extractor.py: NUMERICAL_FEATURE_COLS rozszerzony o NLP_FEATURE_COLS (35 cech lacznie)')
bullet('pipeline.py: galaz tekstowa = NLTKTextPreprocessor -> TfidfVectorizer')

h2('6.4 notebooks/01_eda.ipynb')
para('9 sekcji: rozklad klas, histogramy top-12 cech, korelacja z etykieta, '
     'TF-IDF top words per klasa, analiza URL, analiza NLP, tabela podsumowujaca.')

h2('6.5 Instalacja modelu spaCy')
code('python -m spacy download en_core_web_sm')
doc.add_page_break()

# ---- 7. FAZA 4 ----
h1('7. Faza 4 - API i integracja')

h2('7.1 Nowa struktura agent/')
table_with_header(
    ['Plik', 'Odpowiedzialnosc'],
    [
        ('agent/db.py', 'SQLAlchemy engine, SessionLocal, Prediction ORM (+pole campaign)'),
        ('agent/state.py', 'Wspoldzielony stan modelu ML i pollera (singleton)'),
        ('agent/api/schemas.py', 'Wszystkie modele Pydantic'),
        ('agent/api/routes_predict.py', '/predict, /predict/raw, /model/info'),
        ('agent/api/routes_results.py', '/results, /results/stats, /results/{id}, /results/campaign/{n}'),
        ('agent/api/routes_ingest.py', '/ingest/mailhog, /ingest/campaign/{n}, DELETE /ingest/mailhog/purge'),
        ('agent/ingestion/mailhog_poller.py', 'Async background task - klasyfikacja z MailHog co N sekund'),
        ('agent/main.py', 'Slim: lifespan + router registration + /health + /poller/status'),
    ]
)

h2('7.2 Pelna lista endpointow API')
table_with_header(
    ['Metoda', 'Endpoint', 'Opis'],
    [
        ('GET',    '/health',                  'Stan serwisu + model + poller'),
        ('GET',    '/poller/status',            'Diagnostyka background pollera'),
        ('GET',    '/model/info',               'Typ modelu, sciezka, liczba cech'),
        ('POST',   '/predict',                  'Klasyfikacja z JSON'),
        ('POST',   '/predict/raw',              'Klasyfikacja z surowego MIME (.eml)'),
        ('GET',    '/results',                  'Paginowana historia predykcji'),
        ('GET',    '/results/stats',            'Agregaty: total, rozklad klas'),
        ('GET',    '/results/{id}',             'Szczegoly predykcji'),
        ('GET',    '/results/campaign/{n}',     'Predykcje dla kampanii'),
        ('POST',   '/ingest/mailhog',           'Pobierz i sklasyfikuj z MailHog'),
        ('POST',   '/ingest/campaign/{n}',      'Klasyfikuj .eml z data/raw/{n}/'),
        ('DELETE', '/ingest/mailhog/purge',     'Wyczysc inbox MailHog'),
    ]
)

h2('7.3 MailHog Poller')
para('Asyncio task w lifespan FastAPI. Co 15 sekund (MAILHOG_POLL_INTERVAL) odpytuje MailHog, '
     'identyfikuje nowe wiadomosci po msg_id, klasyfikuje i zapisuje do DB. '
     'Wylaczalny przez MAILHOG_POLLER_ENABLED=false.')
doc.add_page_break()

# ---- 8. STRUKTURA REPO ----
h1('8. Finalna struktura repozytorium')
code(
    'phishing_detector_with_ai_agent/\n'
    '  agent/\n'
    '    api/  schemas.py, routes_predict.py, routes_results.py, routes_ingest.py\n'
    '    features/  extractor.py, nlp_analyzer.py, text_preprocessor.py, url_analyzer.py\n'
    '    ingestion/ email_parser.py, dataset_loader.py, mailhog_poller.py\n'
    '    ml/        pipeline.py\n'
    '    db.py, state.py, main.py\n'
    '  simulation/\n'
    '    gophish_client.py, mailhog_client.py, campaign_generator.py\n'
    '  experiments/\n'
    '    train_model.py\n'
    '  notebooks/\n'
    '    01_eda.ipynb\n'
    '  scripts/\n'
    '    fetch_datasets.py\n'
    '  docker/gophish/config.json\n'
    '  data/{raw,processed,datasets}/\n'
    '  models/\n'
    '  docker-compose.yml, Dockerfile, requirements.txt, .env.example'
)

# ---- 9. ZMIENNE SRODOWISKOWE ----
h1('9. Zmienne srodowiskowe (.env)')
table_with_header(
    ['Zmienna', 'Domyslna', 'Opis'],
    [
        ('DATABASE_URL', 'postgresql+psycopg2://thesis:thesis@db:5432/phishing_lab', 'Polaczenie z PostgreSQL'),
        ('GOPHISH_API_URL', 'http://localhost:3333', 'Adres Gophish Admin API'),
        ('GOPHISH_API_KEY', '(wymagane)', 'Klucz API z panelu Gophish'),
        ('MAILHOG_API_URL', 'http://localhost:8025', 'Adres MailHog API'),
        ('MODEL_PATH', 'models/classifier.joblib', 'Sciezka do modelu'),
        ('MAILHOG_POLLER_ENABLED', 'true', 'Wlacz/wylacz background poller'),
        ('MAILHOG_POLL_INTERVAL', '15', 'Interwal pollera w sekundach'),
        ('PHISHING_THRESHOLD', '0.5', 'Prog decyzyjny klasyfikatora'),
    ]
)

# ---- ZAPIS ----
output = r'C:\Users\filip\PracaDyplomowa\phishing_detector_with_ai_agent\Historia_projektu.docx'
doc.save(output)
print('Zapisano:', output)
